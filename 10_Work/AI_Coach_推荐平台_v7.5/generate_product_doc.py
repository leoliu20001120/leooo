# -*- coding: utf-8 -*-
"""
生成海克斯大乱斗 AI Coach 符文推荐评分机制产品文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ==================== 样式设置 ====================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.5

# 标题样式
for i in range(1, 5):
    h = doc.styles[f'Heading {i}']
    h.font.name = '微软雅黑'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_colored_text(paragraph, text, color=None, bold=False, size=None):
    """为段落添加带颜色的文本"""
    run = paragraph.add_run(text)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    if size:
        run.font.size = size
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return run

def add_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1a1a2e"/>')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.name = '微软雅黑'
        cell._tc.get_or_add_tcPr().append(shading)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = '微软雅黑'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if ri % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f0f4f8"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    return table


# ==================== 文档正文 ====================

# 标题页
title = doc.add_heading('海克斯大乱斗 AI Coach\n符文推荐评分机制产品方案', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_colored_text(p, 'P0版本 · 内部文档', RGBColor(0x6b, 0x72, 0x80), size=Pt(12))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_colored_text(p, '版本：v3.2  |  更新日期：2026年3月31日', RGBColor(0x9c, 0xa3, 0xaf), size=Pt(9))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_colored_text(p, '文档负责人：欧阳婷婷(saintouyang)', RGBColor(0x9c, 0xa3, 0xaf), size=Pt(9))

doc.add_page_break()

# ==================== 目录 ====================
doc.add_heading('目录', level=1)
toc_items = [
    '一、业务背景与产品场景',
    '二、评分函数设计',
    '  2.1 核心公式',
    '  2.2 基础分计算（0-100分）',
    '  2.3 归一化方法与设计依据',
    '  2.4 权重配置（连胜/连败动态调整）',
    '  2.5 英雄胜率纠偏机制',
    '  2.6 UGC极端值处理',
    '  2.7 连胜连败机制',
    '三、黑科技组合加成体系',
    '  3.1 核心规则',
    '  3.2 英雄专属黑科技（最佳拍档）',
    '  3.3 通用黑科技组合（潜力组合）',
    '  3.4 阶段感知加分逻辑',
    '  3.5 套装羁绊加成',
    '四、标签体系与优先级',
    '  4.1 四类标签定义',
    '  4.2 标签判定优先级',
    '  4.3 标签可见性控制',
    '  4.4 强力单卡判定逻辑',
    '  4.5 娱乐标签插入逻辑',
    '五、三分类Logo与自适应阈值',
    '  5.1 Logo分类定义',
    '  5.2 分英雄自适应阈值',
    '  5.3 建议刷新阈值',
    '六、推荐话术生成方案',
    '  6.1 话术核心原则',
    '  6.2 各标签话术来源与生成规则',
    '七、模拟推荐流程',
    '  7.1 阶段与等级说明',
    '  7.2 前置已选卡牌模拟',
    '八、数据来源与知识库',
    '  8.1 数据源一览',
    '  8.2 知识库上传（NPC平台）',
    '九、后续迭代规划',
    '  9.1 二期计划',
    '  9.2 知识图谱构建',
    '十、参数汇总表',
    '附录A：实战案例详解',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.line_spacing = 1.2
    if item.startswith('  '):
        p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ==================== 一、业务背景 ====================
doc.add_heading('一、业务背景与产品场景', level=1)

doc.add_heading('1.1 产品定位', level=2)
doc.add_paragraph(
    '海克斯大乱斗是《英雄联盟》的特色对局模式，玩家在对局中需要在4个阶段分别选择海克斯符文（每次3张同等级随机符文，'
    '每张有1次刷新机会）。AI Coach 是对局内的单向AI助手，通过右侧海克斯机器人在关键时刻主动推送信息，'
    '帮助玩家理解机制、选对符文、感受游戏爽感。'
)
doc.add_paragraph(
    '符文推荐评分机制是 AI Coach 的核心模块，负责为玩家提供：'
)
items = [
    '三分类Logo推荐：👍推荐选取 / 🤔值得考虑 / 🔄建议刷新',
    '标签化推荐：最佳拍档 / 潜力组合 / 强力单卡 / 娱乐（连胜专属）',
    '推荐话术：10-15字的核心推荐理由（区别于卡牌本身描述，提供额外补充信息）',
    '动态评分：根据英雄、阶段、已选符文、连胜/连败状态实时动态调整',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.2 核心设计原则', level=2)
principles = [
    ('降低入门门槛', '帮助新手快速上手，看懂符文、选对符文'),
    ('放大爽感', '连胜时鼓励冒险（推荐黑科技/娱乐），连败时帮稳住局面（推荐高胜率符文）'),
    ('简单可调', '所有参数可在平台上实时调整，方便运营迭代'),
    ('独立评分', '每张符文独立评分，不做强制3选1分布——3张都好都推荐，3张都差都建议刷新'),
]
for title_text, desc in principles:
    p = doc.add_paragraph()
    add_colored_text(p, f'• {title_text}：', bold=True)
    p.add_run(desc)

doc.add_heading('1.3 选符文阶段说明', level=2)
doc.add_paragraph(
    '注意：阶段和符文等级是独立的概念。每个阶段出现的符文等级由游戏随机决定，'
    '并非S1一定出白银、S4一定出棱彩。评分系统中阶段和等级分开处理。'
)
add_table(doc,
    ['阶段', '触发时机', '说明'],
    [
        ['S1', '开局', '优先单卡强度，还谈不上组合'],
        ['S2', '7级', '引入组合判断，可开始凑黑科技'],
        ['S3', '11级', '后期组合成型关键期'],
        ['S4', '15级', '最终选择'],
    ]
)
doc.add_paragraph(
    '符文等级共三种：白银、黄金、棱彩。每个阶段的符文等级由游戏决定，'
    '评分系统根据实际出现的等级进行评分和阈值计算。'
)

doc.add_page_break()

# ==================== 二、评分函数设计 ====================
doc.add_heading('二、评分函数设计', level=1)

doc.add_heading('2.1 核心公式', level=2)
p = doc.add_paragraph()
add_colored_text(p, '最终得分 = (基础分 + 英雄胜率纠偏分) + 黑科技组合加成', bold=True, color=RGBColor(0x1a, 0x1a, 0x2e), size=Pt(12))

doc.add_paragraph(
    '其中：\n'
    '• 基础分（0-100分）= 符文胜率分 × W_winrate + 符文选择率分 × W_pickrate + UGC评分分 × W_ugc\n'
    '• 英雄胜率纠偏分（-5 ~ +8分）= 低胜率英雄提分，高胜率英雄降分\n'
    '• 黑科技组合加成（0 ~ 20分）= 英雄专属/通用组合/套装羁绊加成\n'
    '• 连胜/连败差异通过切换权重配置（标准/连胜/连败三套权重）体现，所有权重均可在参数面板调整'
)

p = doc.add_paragraph()
add_colored_text(p, '💡 关键设计：', bold=True, color=RGBColor(0x22, 0x78, 0xb5))
p.add_run('阶段系数不是独立的乘法因子，而是和黑科技组合加成绑在一起。通用黑科技组合只有S1/S2才给完整加成，'
          'S3/S4需要已选过组合件才给加成；英雄专属黑科技不受阶段限制。')

doc.add_heading('2.2 基础分计算（0-100分）', level=2)
doc.add_paragraph('基础分由三个维度的归一化分数加权求和得到：')

add_table(doc,
    ['维度', '数据源', '归一化范围', '权重(标准)', '说明'],
    [
        ['符文胜率分', '英雄×符文CSV胜率', '0-100分', '0.60', '核心指标，反映符文在该英雄上的实战效果'],
        ['符文选择率分', '英雄×符文CSV选率', '0-100分', '0.15', '社区共识指标，选率高说明玩家认可'],
        ['UGC评分分', '掌盟UGC评分(10分制)', '0-100分', '0.25', '用户偏好修正，包含娱乐/黑科技等主观因素'],
    ]
)

doc.add_heading('2.3 归一化方法与设计依据', level=2)
doc.add_paragraph('三个维度的归一化公式及其设计理由：')

p = doc.add_paragraph()
add_colored_text(p, '① 符文胜率归一化', bold=True)
doc.add_paragraph(
    '公式：score = (胜率 - 45) / (70 - 45) × 100，截断到 [0, 100]\n'
    '• 下界45%：低于45%接近随机水平，设为0分\n'
    '• 上界70%：全局胜率范围约43.93%~67.39%，70%留有余量\n'
    '• 数据来源优先使用英雄×符文胜率（CSV），回退全局胜率（Excel知识库）'
)

p = doc.add_paragraph()
add_colored_text(p, '② 符文选择率归一化', bold=True)
doc.add_paragraph(
    '公式：score = min(选率 / 3.0, 1.0) × 100，截断到 [0, 100]\n'
    '• 饱和点3%：全局选择率范围约0%~2.08%，3%为极高热门符文\n'
    '• 超过3%不再加分，防止"人人都选"的热门符文过度膨胀'
)

p = doc.add_paragraph()
add_colored_text(p, '③ UGC评分归一化（含异常值处理）', bold=True)
doc.add_paragraph(
    '基础公式：score = 调整后评分 / 10 × 100\n'
    '无评分默认50分（中性），详见"2.6 UGC极端值处理"'
)

doc.add_heading('2.4 权重配置（连胜/连败动态调整）', level=2)
doc.add_paragraph('根据玩家连胜/连败状态动态切换权重配置：')

add_table(doc,
    ['场景', '触发条件', 'W_winrate(胜率)', 'W_pickrate(选率)', 'W_ugc(UGC)', '设计理念'],
    [
        ['标准模式', '连胜<3且连败<3', '0.60', '0.15', '0.25', '以胜率为主的均衡推荐'],
        ['连胜模式', '≥3连胜', '0.40', '0.15', '0.45', '降低胜率权重，提高UGC权重，鼓励冒险/娱乐'],
        ['连败模式', '≥3连败', '0.75', '0.15', '0.10', '大幅提高胜率权重，帮助稳住局面'],
    ]
)

p = doc.add_paragraph()
add_colored_text(p, '💡 设计理念：', bold=True, color=RGBColor(0x22, 0x78, 0xb5))
p.add_run('连胜时降低胜率权重、提高UGC权重，鼓励用户追求有趣体验（黑科技/娱乐流）；'
          '连败时大幅提高胜率权重，帮用户稳住局面。三项权重加起来恒等于1.0，简单易调。')

doc.add_heading('2.5 英雄胜率纠偏机制', level=2)
doc.add_paragraph(
    '不同英雄在海克斯大乱斗中胜率差异较大（如火男胜率远高于萃取之神），如果不做纠偏，'
    '低胜率英雄的推荐符文数量会远少于高胜率英雄，导致体验不均。'
)
p = doc.add_paragraph()
add_colored_text(p, '纠偏公式：', bold=True)
doc.add_paragraph(
    'correction = (全英雄平均胜率 - 该英雄胜率) × 纠偏强度系数 × 100 / (上界 - 下界)'
)
add_table(doc,
    ['参数', '值', '说明'],
    [
        ['全英雄平均胜率', '≈49.55%', '来自step1_3真实数据，运行时计算'],
        ['纠偏强度系数', '0.3', '越大纠偏越强'],
        ['纠偏上限', '+8分', '低胜率英雄最多加8分'],
        ['纠偏下限', '-5分', '高胜率英雄最多扣5分'],
    ]
)
doc.add_paragraph(
    '效果：低胜率英雄（如40%）的符文评分整体上移约+5分，使其也有合理数量的推荐符文；'
    '高胜率英雄（如55%）的评分整体下压约-2分，避免推荐过多。'
)

doc.add_heading('2.6 UGC极端值处理', level=2)
doc.add_paragraph(
    'UGC评分来自掌盟APP，数据分布不符合正态分布（长尾+极端值），直接使用会导致：\n'
    '• 极低评分（2.3分）的符文被过度惩罚\n'
    '• 小样本（仅3条评论）的评分波动大，不可靠\n'
    '因此采用两层处理：'
)

p = doc.add_paragraph()
add_colored_text(p, '① 分位数截断（Percentile Clipping）', bold=True)
doc.add_paragraph(
    '• 不依赖正态分布假设，适用于长尾分布\n'
    '• 截断在P5（底部5%）：低于第5百分位的UGC评分统一设为该分位数值\n'
    '• 参数：UGC_CLIP_PERCENTILE = 5.0'
)

p = doc.add_paragraph()
add_colored_text(p, '② 贝叶斯收缩（Bayesian Shrinkage）', bold=True)
doc.add_paragraph(
    '• 解决小样本偏差问题：评分样本越少，越"收缩"到全局均值\n'
    '• 收缩公式：调整后评分 = (样本数 × 原始评分 + 先验权重 × 全局均值) / (样本数 + 先验权重)\n'
    '• 先验权重（UGC_BAYESIAN_PRIOR_WEIGHT）= 30，相当于"虚拟样本数"\n'
    '• 效果：只有3条评论的符文，其UGC评分大幅向全局均值回归；200条评论的符文几乎不受影响'
)

add_table(doc,
    ['处理步骤', '方法', '参数', '作用'],
    [
        ['第1步', '贝叶斯收缩', '先验权重=30', '小样本向全局均值回归，防止小样本偏差'],
        ['第2步', '分位数截断', '截断P5', '极端低值统一截断，防止过度惩罚'],
        ['第3步', '归一化', '/ 10 × 100', '统一到0-100分'],
        ['无评分', '默认50分', '-', '无评分视为中性，不奖不罚'],
    ]
)

doc.add_heading('2.7 连胜连败机制', level=2)
doc.add_paragraph(
    '连胜/连败差异通过切换权重配置（三套权重profile）来体现：'
)
add_table(doc,
    ['场景', '触发条件', 'W_winrate', 'W_pickrate', 'W_ugc', '设计理念'],
    [
        ['标准模式', '连胜<3且连败<3', '0.60', '0.15', '0.25', '以胜率为主的均衡推荐'],
        ['连胜模式', '≥3连胜', '0.40', '0.15', '0.45', '降低胜率权重，提高UGC权重，鼓励冒险/娱乐'],
        ['连败模式', '≥3连败', '0.75', '0.15', '0.10', '大幅提高胜率权重，帮助稳住局面'],
    ]
)
p = doc.add_paragraph()
add_colored_text(p, '💡 ', bold=True, color=RGBColor(0x22, 0x78, 0xb5))
p.add_run('三项权重加起来恒等于1.0，简单易调。所有权重均可在参数面板实时调整。')

doc.add_page_break()

# ==================== 三、黑科技组合加成体系 ====================
doc.add_heading('三、黑科技组合加成体系', level=1)

doc.add_heading('3.1 核心规则', level=2)
doc.add_paragraph(
    '黑科技组合加成是评分函数中的额外加分项，最终封顶20分（BLACKTECH_BONUS_CAP = 20）。'
    '判定需要考虑三个维度：组合是否适配当前英雄、当前阶段、是否已选过组合件。'
)

add_table(doc,
    ['场景', '标签', '加成', '说明'],
    [
        ['英雄专属黑科技', '最佳拍档', '+20', '如火男+炼狱导管，单卡即生效，不受阶段限制'],
        ['通用组合成型+适配英雄', '潜力组合', '+25', '已选过组合另一半+适配英雄→组合成型！'],
        ['通用组合成型+不适配英雄', '-', '+5', '组合效果在但英雄用不好'],
        ['S1/S2+通用组合适配+未成型', '潜力组合', '+10', '前期有机会凑齐，值得期待'],
        ['S3/S4+适配但未成型', '-', '+3', '晚期凑齐概率低，小幅鼓励'],
        ['不适配英雄+未成型', '-', '+0', '直接跳过，不加分'],
        ['套装羁绊2件及以上', '-', '+5×件数', '独立于黑科技标签，凑够2件就加分'],
    ]
)

p = doc.add_paragraph()
add_colored_text(p, '⚠️ 加成总封顶：', bold=True, color=RGBColor(0xea, 0x38, 0x38))
p.add_run('所有黑科技加成（专属+组合+套装）累加后封顶30分（BlacktechMatcher返回上限），'
          '进入评分函数时再封顶20分（BLACKTECH_BONUS_CAP），防止多组合叠加导致分数爆炸。')

doc.add_heading('3.2 英雄专属黑科技（最佳拍档标签）', level=2)
doc.add_paragraph(
    '来源：黑科技组合分析_v5.xlsx「英雄专属黑科技」sheet，共985条全量英雄×符文匹配记录。\n'
    '原始数据来自arammayhem.com第三方攻略中标记为"黑科技"或"神级"的英雄+符文特殊交互。'
)
doc.add_paragraph('代表性案例：')
add_table(doc,
    ['英雄', '符文', '评级', '推荐话术', '黑科技原因'],
    [
        ['火男', '炼狱导管', '神级(33)', '单技能叠四层', '被动灼烧叠加，唯一真神符文'],
        ['法外狂徒', '亮出你的剑', '神级(20)', '一枪一个', '大招+符文爆发一击必杀'],
        ['暗夜猎手', '双刀流', '神级(15)', '快速三环', '攻速叠加被动银弩三环'],
        ['布隆', '连拨击锤', '黑科技', '一下满被动', 'A一下直接满被动眩晕'],
        ['腕豪', '量子计算', '黑科技', '四倍真伤', '大招期间4倍内圈伤害'],
    ]
)
doc.add_paragraph(
    '特殊说明：英雄专属黑科技不受阶段限制（S1-S4均生效），因为是"单卡即生效"的强交互，'
    '不需要凑组合件。标签显示为"最佳拍档"。'
)

doc.add_heading('3.3 通用黑科技组合（潜力组合标签）', level=2)
doc.add_paragraph('25个2符文核心组合，每个至少10个英雄验证。以下为代表性组合：')
add_table(doc,
    ['#', '流派', '符文1', '符文2', '推荐话术', '机制协同', '适配英雄类型'],
    [
        ['1', 'AP风筝流', '缩小引擎', '超凡邪恶', '小身板高法强', '移速风筝+无限叠法强', 'AP法师'],
        ['2', '控制叠层流', '扇巴掌', '缩小引擎', '控住缩小碾压', '控制叠适应力+变小获急速', '控制型法师'],
        ['3', '灼烧流', '炼狱导管', '祖母的辣椒油', '双灼烧叠加', '双灼烧源使伤害翻倍', 'DOT法师'],
        ['4', '阵亡爆炸流', '俯冲轰炸', '小丑学院', '死了也炸人', '双重阵亡惩罚', '突进战士/刺客'],
        ['5', '体型流', '坦克引擎', '任务：钢化你心', '叠血变大', '双叠血+体型增长', '坦克/战士'],
        ['...', '...', '...', '...', '...', '（共25个组合）', '...'],
    ]
)

doc.add_heading('3.4 阶段感知加分逻辑', level=2)
doc.add_paragraph(
    '黑科技组合的加分与阶段强绑定，核心逻辑：'
)
items = [
    'S1/S2阶段：通用组合适配英雄时给潜力组合标签 + 加分（此时还有机会凑齐组合件）',
    'S3/S4阶段：通用组合如果没有前面阶段已选过另一半，仅给小幅加分但不给潜力组合标签',
    'S3/S4阶段：如果前面阶段已选过组合另一半，视为"组合成型"，给最高加成+标签',
    '英雄专属黑科技：任何阶段（S1-S4）均给加成和标签，不受阶段限制',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('阶段感知判定详细规则：')
add_table(doc,
    ['阶段', '组合状态', '适配英雄', '加成', '标签', '话术'],
    [
        ['任意', '英雄专属黑科技', '-', '+20', '最佳拍档', '专属话术'],
        ['任意', '已选另一半+适配', '✅', '+25', '潜力组合', '组合成型！+流派话术'],
        ['任意', '已选另一半+不适配', '❌', '+5', '-', '-'],
        ['S1/S2', '未选另一半+适配', '✅', '+10', '潜力组合', '流派话术'],
        ['S1/S2', '未选另一半+不适配', '❌', '+0', '-', '跳过'],
        ['S3/S4', '未选另一半+适配', '✅', '+3', '-', '晚期适配'],
        ['S3/S4', '未选另一半+不适配', '❌', '+0', '-', '跳过'],
    ]
)

doc.add_heading('3.5 套装羁绊加成', level=2)
doc.add_paragraph(
    '游戏内9个官方套装羁绊系统，独立于黑科技标签。凑够2件及以上给加分：\n'
    '公式：套装加成 = 5 × 套装内已有件数\n'
    '示例：叠角龙套装已有3件 → +15分\n\n'
    '9个套装：叠角龙(9符文)、神龙烈焰(5符文)、完全自动化(8符文)、金币雨(7符文)、'
    '俯冲炸弹(4符文)、下雪天(5符文)、喂呜喂呜(7符文)、掷骰狂人(4符文)、大法师(5符文)'
)

doc.add_page_break()

# ==================== 四、标签体系 ====================
doc.add_heading('四、标签体系与优先级', level=1)

doc.add_heading('4.1 四类标签定义', level=2)
add_table(doc,
    ['标签', '含义', '触发条件', '显示条件'],
    [
        ['最佳拍档', '英雄专属黑科技，单卡即生效', '英雄+符文在专属黑科技库中', '推荐选取Logo下始终显示'],
        ['潜力组合', '通用黑科技组合匹配', '符文在某组合中且适配当前英雄', '推荐选取Logo下始终显示'],
        ['强力单卡', '英雄×符文胜率TOP15%', '符文在该英雄所有符文胜率中排前15%', '推荐选取Logo下且无上两类标签时补充'],
        ['娱乐', '连胜专属趣味符文', '符文在娱乐符文池中 + 连胜≥3', '仅连胜模式下显示，且该英雄推荐中无足够前三类标签'],
    ]
)

doc.add_heading('4.2 标签判定优先级', level=2)
doc.add_paragraph('标签判定按以下优先级进行，同一符文最多获得1个标签：')
add_table(doc,
    ['优先级', '标签', '判定逻辑', '说明'],
    [
        ['1（最高）', '最佳拍档', '英雄专属黑科技匹配', '走BlacktechMatcher.match()中hero_exclusive类型'],
        ['2', '潜力组合', '通用组合匹配（combo_complete/combo_potential_fit等）', '需组合适配英雄+阶段规则'],
        ['3', '娱乐', '符文名在"娱乐符文"sheet中', '黑科技组合分析_v5.xlsx加载的娱乐符文池'],
        ['4', '强力单卡', '推迟判定（在标签可见性阶段补充）', '需要知道推荐状态后才能判定'],
        ['5', '无标签', '以上均不满足', '普通符文'],
    ]
)

doc.add_heading('4.3 标签可见性控制', level=2)
doc.add_paragraph('并非所有标签都会对玩家展示，可见性由以下规则控制：')
items = [
    '仅"推荐选取"Logo的符文才展示标签（值得考虑和建议刷新的不展示标签）',
    '标准模式下：娱乐标签不展示（visible_tag = None）',
    '连胜模式下：娱乐标签正常展示',
    '强力单卡标签补充：推荐选取中如果最佳拍档+潜力组合不足3个，从无标签的推荐符文中补充强力单卡（最多补2个）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.4 强力单卡判定逻辑', level=2)
doc.add_paragraph(
    '强力单卡不在初始标签判定中产生，而是在标签可见性阶段作为补充：\n\n'
    '判定条件：该英雄×该符文的胜率在该英雄所有符文胜率中排名前15%\n'
    '具体实现：'
)
items = [
    '获取该英雄在champion_augment_stats中的所有符文胜率',
    '使用numpy.percentile计算P85（即TOP15%的阈值）',
    '当前符文胜率 ≥ P85 → 判定为强力单卡',
    '参数：STRONG_CARD_TOP_PERCENT = 15.0',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.5 娱乐标签插入逻辑', level=2)
doc.add_paragraph(
    '娱乐标签专为连胜玩家设计，核心逻辑是"连胜时鼓励冒险"：'
)

p = doc.add_paragraph()
add_colored_text(p, '触发条件：', bold=True)
p.add_run('连胜 ≥ 3 时激活娱乐逻辑')

p = doc.add_paragraph()
add_colored_text(p, '娱乐符文池来源：', bold=True)
p.add_run('黑科技组合分析_v5.xlsx「娱乐符文」sheet中的全量符文')

p = doc.add_paragraph()
add_colored_text(p, '连胜娱乐调整流程（分数调整法）：', bold=True)

items = [
    '找出所有"最佳拍档"和"强力单卡"标签的符文',
    '各按分数排序，后50%（WINNING_DEMOTE_PERCENT = 50%）的卡片施加负分数惩罚 → 自然跌落到"值得考虑"',
    '找出所有"娱乐"标签的符文，施加+15分奖励（ENTERTAINMENT_BOOST = 15） → 自然提升到"推荐选取"',
    '重新排序和重新判定Logo',
    '标签可见性：连胜模式下娱乐标签正常展示；标准模式下隐藏',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
add_colored_text(p, '💡 核心目标：', bold=True, color=RGBColor(0x22, 0x78, 0xb5))
p.add_run('让连胜玩家的推荐列表中出现更多"有趣但不一定最强"的符文，放大游戏乐趣。'
          '同时将实用性较低的最佳拍档/强力单卡降级到"值得考虑"，保持推荐池质量。')

doc.add_page_break()

# ==================== 五、三分类Logo ====================
doc.add_heading('五、三分类Logo与自适应阈值', level=1)

doc.add_heading('5.1 Logo分类定义', level=2)
add_table(doc,
    ['Logo', '图标', '颜色', '含义', '触发条件'],
    [
        ['推荐选取', '👍', '绿色 #22c55e', '强力推荐该符文', '最终得分 ≥ 推荐阈值'],
        ['值得考虑', '🤔', '黄色 #eab308', '可以考虑选择', '得分 ≥ 刷新阈值 且 < 推荐阈值'],
        ['建议刷新', '🔄', '灰色 #9ca3af', '建议使用刷新机会', '得分 < 刷新阈值'],
    ]
)

doc.add_heading('5.2 分英雄自适应阈值', level=2)
doc.add_paragraph(
    '为解决不同英雄胜率差异导致的推荐数量不均，采用分英雄×分等级的自适应阈值：'
)
items = [
    '对每个英雄×每个等级的所有符文评分，按分数降序排列',
    '取第TARGET_RECOMMEND_PER_LEVEL(=4)名的分数作为推荐阈值',
    '取第(TARGET_RECOMMEND_PER_LEVEL × 2)(=8)名的分数作为考虑阈值',
    '确保推荐阈值 > 考虑阈值，否则考虑阈值 = 推荐阈值 - 5',
    '目标：每英雄每等级约4个推荐（最少2个、最多6个），总推荐约12个',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('无英雄信息时使用固定阈值兜底：')
add_table(doc,
    ['符文等级', '推荐阈值', '考虑阈值'],
    [
        ['白银', '42', '28'],
        ['黄金', '45', '30'],
        ['棱彩', '35', '22'],
    ]
)

doc.add_heading('5.3 建议刷新阈值', level=2)
doc.add_paragraph(
    '建议刷新阈值 = 排名底部20%（REFRESH_BOTTOM_PERCENT = 20.0）的分数。\n'
    '即：对该英雄×该等级的所有评分符文，排名在最后20%的符文建议刷新。\n'
    '中间区域（推荐阈值 ~ 刷新阈值之间）= 值得考虑。'
)

doc.add_page_break()

# ==================== 六、推荐话术生成方案 ====================
doc.add_heading('六、推荐话术生成方案', level=1)

doc.add_heading('6.1 话术核心原则', level=2)
items = [
    '推荐话术 ≤ 15字（最好控制在10字左右），极短标签描述核心机制',
    '不出现具体胜率/选取率数字（游戏内已有数据展示）',
    '区别于符文卡牌本身描述（卡牌描述游戏里已有显示），提供玩家额外补充信息',
    '用定性描述：如"表现极为出色"、"社区公认强力"',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2 各标签话术来源与生成规则', level=2)
add_table(doc,
    ['标签类型', '话术来源', '生成方式', '示例'],
    [
        ['最佳拍档\n(英雄专属黑科技)', '黑科技组合分析V5中的"黑科技原因"', 'AI润色加工（15字以内、10字左右最佳）\n基于英雄专属黑科技的原因进行润色', '火男+炼狱导管→"单技能叠四层"\n布隆+连拨击锤→"一下满被动"'],
        ['潜力组合\n(通用黑科技)', '通用黑科技组合的"机制协同"描述', 'AI润色（基于为什么协同/机制进行润色）', 'AP风筝流→"小身板高法强"\n灼烧流→"双灼烧叠加"'],
        ['强力单卡', '知识库中的第三方描述+玩家补充描述', 'AI润色（基于arammayhem攻略描述\n和UGC补充评论）', '"法师通用神器"\n"坦度拉满"'],
        ['娱乐', 'UGC评论', '主要基于掌盟UGC热门评论润色', '"无限雪球"\n"超远击飞"'],
        ['值得考虑\n(无标签)', '卡牌官方描述（掌盟）', '简要生成推荐理由', '"可以一试"\n"数据优秀"'],
        ['建议刷新', '-', '委婉劝玩家刷新', '"可能不太适合法师哦"\n"建议看看刷新"'],
    ]
)

p = doc.add_paragraph()
add_colored_text(p, '⚠️ 重要说明：', bold=True, color=RGBColor(0xea, 0x38, 0x38))
p.add_run('所有推荐话术都必须区别于符文卡牌本身的描述！因为卡牌描述在游戏里已有显示，'
          '我们需要给玩家额外的补充信息——告诉他们"为什么要选"以及"选了之后能获得什么"。')

doc.add_page_break()

# ==================== 七、模拟推荐流程 ====================
doc.add_heading('七、模拟推荐流程', level=1)

doc.add_heading('7.1 阶段与等级映射', level=2)
doc.add_paragraph('模拟推荐面板的操作流程：连胜/连败 → 英雄 → 阶段 → 等级（自动联动）')
add_table(doc,
    ['阶段', '对应等级', '是否需要前置卡牌模拟'],
    [
        ['S1', '白银', '否（第一次选卡）'],
        ['S2', '黄金', '是（模拟S1已选的卡牌）'],
        ['S3', '黄金', '是（模拟S1+S2已选的卡牌）'],
        ['S4', '棱彩', '是（模拟S1+S2+S3已选的卡牌）'],
    ]
)

doc.add_heading('7.2 前置已选卡牌模拟', level=2)
doc.add_paragraph(
    'S2/S3/S4阶段的模拟推荐需要考虑前面阶段玩家已选的符文，因为黑科技组合加分依赖于"是否已选过组合件"。\n\n'
    '模拟方式：'
)
items = [
    '手动添加：从下拉列表中选择前面阶段已选的符文',
    '随机模拟：系统从黑科技组合的符文池中随机抽取（每个前置阶段1张），模拟真实场景',
    '清空：清除所有前置卡牌',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    '传递给后端的参数：selected_augments（前置已选符文列表），BlacktechMatcher.match()会根据'
    'selected_augments判断组合是否已成型，决定加分多少。'
)

doc.add_page_break()

# ==================== 八、数据来源与知识库 ====================
doc.add_heading('八、数据来源与知识库', level=1)

doc.add_heading('8.1 数据源一览', level=2)
add_table(doc,
    ['数据类型', '文件/来源', '关键字段', '用途'],
    [
        ['SQL取数-符文全局数据', 'step1_1_augment_stats.csv', 'win_rate, show_rate', '全局符文胜率&选取率（兜底）'],
        ['SQL取数-英雄×符文', 'step1_2_champion_augment_stats.csv', 'win_rate, show_rate', '英雄×符文胜率（核心数据源）'],
        ['SQL取数-英雄出场率', 'step1_3_champion_pick_rate.csv', 'pick_rate, win_rate', '英雄胜率纠偏基准'],
        ['SQL取数-符文组合', 'step1_4_pair_stats.csv', 'win_rate', '2符文组合胜率验证'],
        ['SQL取数-英雄×符文组合', 'step1_5_champion_pair_stats.csv', 'win_rate', '英雄×符文组合胜率'],
        ['Excel知识库', '海克斯大乱斗符文知识库_v5.xlsx', '符文信息+等级+UGC', '符文基础信息+套装分组'],
        ['Excel知识库', '黑科技组合分析_v5.xlsx', '通用组合/专属/娱乐', '黑科技组合+标签判定'],
        ['JSON知识库', 'zhangmeng_ugc.json', 'score, hot_comments', 'UGC评分+热门评论'],
        ['JSON知识库', 'hextech_synergies.json', 'augments, tier_effects', '9个官方套装'],
        ['JSON知识库', 'augment_id_map.json', '符文ID↔中文名', '符文ID映射'],
    ]
)

doc.add_heading('8.2 知识库上传（NPC平台）', level=2)
doc.add_paragraph(
    '符文推荐的知识库需要上传至NPC平台，以便线上服务使用：'
)
items = [
    '离线预计算：符文评分阈值、标签、预选池均为离线计算后上传',
    '热更机制：知识库内容支持热更新（不需要重新合包），可在15号前完成',
    '需雨琪团队支持上传流程',
    '后续迭代可能迁移至云端支持实时数据',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ==================== 九、后续迭代规划 ====================
doc.add_heading('九、后续迭代规划', level=1)

doc.add_heading('9.1 二期计划', level=2)
add_table(doc,
    ['模块', '优化方向', '具体内容', '优先级'],
    [
        ['实时状态', '加入对局实时状态变量', '连胜/连败、KDA、经济水平作为评分变量\n实时通过OMG接口获取帧率级数据', 'P1'],
        ['推荐系统', '在线推理 + 云端服务', '从本地离线迁移到云端实时计算\n支持实时数据流上报和推荐', 'P1'],
        ['标签体系', '多标签优先级 + 知识图谱', '构建英雄-符文边关系的知识网络\n在线推理时按边检索标签', 'P1'],
        ['数据回收', '埋点 + 效果验证', 'TRACE ID跟踪单局选择\n对比推荐结果与实际选择\n计算采纳率', 'P1'],
        ['推荐话术', 'AI生成 + 质量审核', '使用AI Agent批量生成推荐理由\n人工审核TOP100高频符文话术', 'P2'],
        ['用户画像', '个性化推荐', '基于用户历史对局数据做协同过滤\n不同段位倾向性推荐', 'P2'],
        ['装备联动', '符文→装备推荐', '基于已选符文推荐适配装备\n建立"英雄×符文→装备→胜率"对照表', 'P1'],
    ]
)

doc.add_heading('9.2 知识图谱构建', level=2)
doc.add_paragraph(
    '当前系统的标签是"离散"的（每轮选符文时独立判定），无法教会玩家决策链路。'
    '二期计划构建知识图谱：'
)
items = [
    '节点：英雄节点 + 符文节点',
    '边：英雄-符文之间的关系（最佳拍档/潜力组合/强力单卡/娱乐等）',
    '边属性：胜率、选率、组合协同描述、话术等',
    '推理方式：在线检索——知道当前英雄和已选符文，扫描候选符文节点的边关系',
    '优势：标签从离散变成网络化，支持跨阶段的组合路径推荐',
    '落地方式：可使用图数据库或向量检索实现',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    '⚡ 知识图谱的核心价值：从"授人以鱼"（直接告诉选什么）升级为"授人以渔"（教会玩家为什么这么选），'
    '构建玩家的决策认知链路。同时大幅降低运营维护成本——图谱更新后推理自动适配，不需要人为逐条维护规则。'
)

doc.add_page_break()

# ==================== 十、参数汇总表 ====================
doc.add_heading('十、参数汇总表', level=1)
doc.add_paragraph('以下为P0版本所有可调参数的完整列表，均可在推荐平台上实时调整：')

add_table(doc,
    ['参数名', '当前值', '分类', '说明'],
    [
        ['W_winrate', '0.60', '权重', '标准模式胜率权重'],
        ['W_pickrate', '0.15', '权重', '标准模式选率权重'],
        ['W_ugc', '0.25', '权重', '标准模式UGC权重'],
        ['WR_FLOOR', '45.0', '归一化', '胜率归一化下界'],
        ['WR_CEILING', '70.0', '归一化', '胜率归一化上界'],
        ['PR_SATURATION', '3.0', '归一化', '选率饱和点'],
        ['UGC_MAX', '10.0', '归一化', 'UGC评分最大值'],
        ['BLACKTECH_BONUS_CAP', '20', '黑科技', '黑科技加成封顶'],
        ['HERO_CORRECTION_STRENGTH', '0.3', '英雄纠偏', '纠偏强度系数'],
        ['HERO_CORRECTION_MAX', '+8.0', '英雄纠偏', '纠偏上限'],
        ['HERO_CORRECTION_MIN', '-5.0', '英雄纠偏', '纠偏下限'],
        ['TARGET_RECOMMEND_PER_LEVEL', '4', '阈值', '每等级目标推荐数'],
        ['MIN_RECOMMEND_PER_LEVEL', '2', '阈值', '每等级最少推荐'],
        ['MAX_RECOMMEND_PER_LEVEL', '6', '阈值', '每等级最多推荐'],
        ['STRONG_CARD_TOP_PERCENT', '15.0', '标签', '强力单卡TOP百分比'],
        ['UGC_CLIP_PERCENTILE', '5.0', 'UGC处理', '分位数截断P5'],
        ['UGC_BAYESIAN_PRIOR_WEIGHT', '30', 'UGC处理', '贝叶斯先验权重'],
        ['REFRESH_BOTTOM_PERCENT', '20.0', '阈值', '建议刷新底部百分比'],
        ['WINNING_DEMOTE_PERCENT', '50.0', '连胜', '连胜降级比例'],
        ['ENTERTAINMENT_BOOST', '15.0', '连胜', '娱乐符文加分'],
    ]
)

doc.add_page_break()

# ==================== 附录A ====================
doc.add_heading('附录A：实战案例详解', level=1)

doc.add_heading('案例1：丽桑卓S1 + 缩小引擎', level=2)
doc.add_paragraph('英雄: 丽桑卓, 阶段: S1, 当前符文: 缩小引擎, 已选符文: 无')
doc.add_paragraph('判定过程：')
items = [
    '英雄专属黑科技：丽桑卓+缩小引擎 → 不在专属库 → 跳过',
    '通用黑科技组合：找到3个适配组合（AP风筝流+10、控制叠层流+10、坦控叠层流+10）',
    '不适配的组合直接跳过（如灼烧缩小流，丽桑卓不在适配英雄列表）',
    '套装：单卡不触发 → +0',
    '结果：加成+20（封顶），标签=潜力组合，话术="小身板高法强"',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('案例2：丽桑卓S3 + 缩小引擎（S1已选超凡邪恶）', level=2)
doc.add_paragraph('英雄: 丽桑卓, 阶段: S3, 当前符文: 缩小引擎, 已选符文: [超凡邪恶]')
doc.add_paragraph('判定过程：')
items = [
    '缩小引擎+超凡邪恶（AP风筝流）→ 已有另一半+丽桑卓适配 → 组合成型！+25',
    '其余适配组合 S3+未成型 → 各+3',
    '叠角龙套装2件 → +10',
    '结果：加成+20（封顶），标签=潜力组合，话术="组合成型！小身板高法强"',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('案例3：火男 任何阶段 + 炼狱导管', level=2)
doc.add_paragraph('英雄: 火男, 阶段: 任意, 当前符文: 炼狱导管')
doc.add_paragraph('判定过程：')
items = [
    '英雄专属黑科技：火男+炼狱导管 → 在专属库中（神级score:33）→ +20',
    '标签=最佳拍档，话术="单技能叠四层"',
    '不受阶段限制（S1到S4都是最佳拍档）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('案例4：奥拉夫S1 + 缩小引擎', level=2)
doc.add_paragraph('英雄: 奥拉夫, 阶段: S1, 当前符文: 缩小引擎, 已选符文: 无')
doc.add_paragraph('判定过程：')
items = [
    '通用组合：缩小引擎参与4个组合，但奥拉夫是物理战士，全部不适配 → 全部跳过，+0',
    '结果：加成+0，无特殊标签',
    '说明：不适配英雄的组合 = 对该英雄没意义 = 不加分',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('案例5：连胜模式下的娱乐标签插入', level=2)
doc.add_paragraph('英雄: 任意, 连胜: 3连胜, 等级: 黄金')
doc.add_paragraph('娱乐逻辑触发过程：')
items = [
    '找出所有最佳拍档和强力单卡的推荐符文',
    '按分数排序，后50%施加负分数惩罚 → 降级到"值得考虑"',
    '找出所有在娱乐符文池中的符文，施加+15分奖励 → 提升到"推荐选取"',
    '重新排序和判定Logo',
    '娱乐标签可见（连胜模式下visible_tag = "娱乐"）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ==================== 保存 ====================
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "产品文档",
    "海克斯大乱斗AI Coach符文推荐评分机制产品方案.docx"
)
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"文档已保存到: {output_path}")
